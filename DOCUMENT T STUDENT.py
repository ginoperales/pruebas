from docx import Document

# Crear un documento Word
doc = Document()

# Título del documento
doc.add_heading("Uso de la Distribución t de Student: Una Cola vs. Dos Colas", level=1)

# Introducción general
doc.add_paragraph(
    "En el marco de una tesis, el uso de la distribución t de Student de una cola o de dos depende del "
    "contexto del análisis estadístico y de la hipótesis planteada. A continuación, se describen las situaciones "
    "en las que se utiliza cada una."
)

# Sección 1: Prueba de una cola
doc.add_heading("1. Prueba de una cola", level=2)
doc.add_paragraph(
    "Cuándo se usa:\n"
    "- Cuando la hipótesis alternativa (H₁) es unidireccional, es decir, esperas un efecto o diferencia en una dirección específica.\n"
    "- Ejemplo: Estás evaluando si la resistencia de un nuevo material es mayor que la resistencia de un material convencional."
)
doc.add_paragraph(
    "Hipótesis:\n"
    "- H₀: La resistencia del nuevo material es menor o igual a la del convencional (μ ≤ μ₀).\n"
    "- H₁: La resistencia del nuevo material es mayor que la del convencional (μ > μ₀)."
)
doc.add_paragraph("Zona crítica: Está en un extremo de la distribución.")

# Sección 2: Prueba de dos colas
doc.add_heading("2. Prueba de dos colas", level=2)
doc.add_paragraph(
    "Cuándo se usa:\n"
    "- Cuando la hipótesis alternativa es bidireccional, es decir, no sabes en qué dirección podría estar el efecto o la diferencia, "
    "pero esperas que exista una diferencia significativa.\n"
    "- Ejemplo: Estás evaluando si un nuevo tratamiento médico tiene un efecto diferente (mejor o peor) que el tratamiento actual."
)
doc.add_paragraph(
    "Hipótesis:\n"
    "- H₀: El efecto del nuevo tratamiento es igual al del tratamiento actual (μ = μ₀).\n"
    "- H₁: El efecto del nuevo tratamiento es diferente del tratamiento actual (μ ≠ μ₀)."
)
doc.add_paragraph("Zona crítica: Está en ambos extremos de la distribución.")

# Ejemplo práctico en una tesis
doc.add_heading("Ejemplo práctico en una tesis", level=2)
doc.add_paragraph(
    "Si en tu tesis estás analizando los resultados de un experimento con un nuevo diseño estructural:\n"
    "- Prueba de una cola: Si quieres probar si el nuevo diseño soporta más carga que el diseño tradicional.\n"
    "- Prueba de dos colas: Si solo quieres probar si hay una diferencia significativa en la capacidad de carga entre "
    "el nuevo diseño y el tradicional, sin importar si es mayor o menor."
)

# Consideraciones clave
doc.add_heading("Consideraciones clave", level=2)
doc.add_paragraph(
    "1. Plantea las hipótesis correctamente: Define si esperas una mejora específica (una cola) o cualquier diferencia (dos colas).\n"
    "2. Nivel de significancia (α):\n"
    "- En pruebas de dos colas, α se divide entre las dos colas.\n"
    "- En pruebas de una cola, todo α se concentra en un lado de la distribución.\n"
    "3. Justifica tu elección: En el marco de la tesis, incluye una explicación de por qué escogiste una cola o dos, basado en el diseño "
    "del experimento y tus objetivos de investigación."
)

# Guardar el documento
file_path = "C:\\Users\\gin_0\\Downloads\\PRUEBAS\\Uso_de_Distribucion_t_Student.docx" #TAMBIEN PUEDES UTILIZAR LOS SEPARADORES / EN REEEMPLADO DE \\
file_path
